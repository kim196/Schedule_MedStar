import os
import re
from datetime import datetime
import pandas as pd
from docx import Document

# parses sheet into document per day
# output / input redirection---------------------
EXCEL_PATH = "Schedule.xlsx"
TEMPLATE_PATH = "NPPA Assignments.docx"
OUTPUT_DIR = "Created_Docs" # outputs new doc in this directory
os.makedirs(OUTPUT_DIR, exist_ok=True)

# shift filters ------------ 
def is_valid_shift(shift):
    if not isinstance(shift, str):
        return False
    shift = shift.upper().strip()
    # if any of these are included in a shift, then skip and do not include in the NPPA Assignment
    if any(bad in shift for bad in ["FSH", "PTO", "SH", "OT"]):
        return False
    # use regex to find correct formatted time -> \d+ = 1 or more digits , then 0 or more spaces, - , then 1 or more digits
    # but one thing possible instead of having rose fix all the dates, change the regex to also just except digits (ex 630) no dashes needed represents 6:30- 6:30
   
    # return bool(re.search(r"\d+\s*-\s*\d+", shift))
    return bool(re.search(r"\d+(\s*-\s*\d+)?", shift))


# checks whether a given shift qualifies someone to be assigned as "Point" in your NPPA assignment sheet.
def is_point(shift):
    return isinstance(shift, str) and shift.strip().upper().endswith("P")


def is_late_shift(shift):
    # if not isinstance(shift, str):
    #   return False
    # shift = shift.strip()
    # return shift.endswith("-9") or shift.endswith("9") or "9" in shift.split("-")[-1]

    # returns true if 9 appears anywhere in shift
    if not isinstance(shift, str):
        return False
    return "9" in shift

# Format bold and labels
def add_bold_labels(paragraph, label, name, pager_time):
    paragraph.clear()
    if label:
        paragraph.add_run(label + " ").bold = True
    paragraph.add_run(name)
    paragraph.add_run(" Pager/Time: ").bold = True
    paragraph.add_run(pager_time)

# Date Parse-----------
def extract_daily_assignments(sheet, target_date_str):
    target_date = datetime.strptime(target_date_str, "%Y-%m-%d").date()
    assignments = []

    # 17 rows per day (17 rotations of people)
    for block_start in range(0, len(sheet), 17):
        if block_start + 2 >= len(sheet):
            break
        date_row = sheet.iloc[block_start + 1]
        name_rows = sheet.iloc[block_start + 2:block_start + 15]

        for col_idx, cell in enumerate(date_row):
            if isinstance(cell, datetime) and cell.date() == target_date:
                for _, row in name_rows.iterrows():
                    name = row[0]
                    time = row[col_idx]
                    if pd.notna(time) and isinstance(name, str) and ("NP" in name or "PA" in name):
                        assignments.append((name.strip(), str(time).strip()))
                return assignments
    return []

# Fill out Docx-------------
def fill_half_sheet(date_str, assignments):
    doc = Document(TEMPLATE_PATH)

    valid = [(n, t) for (n, t) in assignments if is_valid_shift(t)]
    point = next(((n, t) for (n, t) in valid if is_point(t)), None)
    late = next(((n, t) for (n, t) in valid if is_late_shift(t)), None)
    dr_wang = next(((n, t) for (n, t) in valid if "JOHN HERMANN" in n.upper()), None)

    others = [
        (n, t) for (n, t) in valid
        if (dr_wang is None or n != dr_wang[0]) and
           (point is None or n != point[0]) and
           (late is None or n != late[0])
    ][:4]

    for para in doc.paragraphs:
        text = para.text

        if "Date:" in text:
            para.text = f"NP/PA Assignments                      Date: {date_str}"

        elif "Point:" in text and point:
            add_bold_labels(para, "Point:", point[0], point[1])

        elif "Dr. Wang:" in text and dr_wang:
            add_bold_labels(para, "Dr. Wang:", dr_wang[0], dr_wang[1])

        elif "Pager/Time:_____________" in text and "Dr. Wang" not in text and "Point" not in text and "9 PM" not in text:
            if others:
                n, t = others.pop(0)
                add_bold_labels(para, "", n, t)

        elif "9 PM/Late:" in text and late:
            add_bold_labels(para, "9 PM/Late:", late[0], late[1])

    save_path = os.path.join(OUTPUT_DIR, f"NPPA_Assignments_{date_str}.docx")
    doc.save(save_path)
    print(f" Saved: {save_path}")
    return save_path

# MAIN---------------------
def main():
    date_input = input("Enter a date (YYYY-MM-DD): ").strip()
    try:
        datetime.strptime(date_input, "%Y-%m-%d")
    except ValueError:
        print("Invalid date format. Use YYYY-MM-DD.")
        return

    xl = pd.ExcelFile(EXCEL_PATH)
    for sheet_name in xl.sheet_names:
        sheet = xl.parse(sheet_name, header=None)
        assignments = extract_daily_assignments(sheet, date_input)
        if assignments:
            fill_half_sheet(date_input, assignments)
            return
    print("No assignments found for that date.")

if __name__ == "__main__":
    main()

# P and W must be assigned once to each person, only through M-F, cannot have P and W in the same time 
# P and W must be assigned to someone working before 12 pm (not including the 12 so 11 and before am). This must be equal times within 6 weeks
# if anyone works 9-9pm (so considered late) cannot assigned to P (Point) but can be assigned to W. P and W must be equally distributed 

# Ex: 2025-07-23 Still having an issue that a single digit like "6" should be read as "6-6" but it is not being read because its not formatted right
# Option 1 : Manually fix all other shifts to be # - # or 
# Option 2 : Code so it will count a singular digit 

# im struggling to do this so i think easiest way is to scan schedule, fill it out and return a new sheets with equal P and W distribution, 
# then for testing we can count up P and W